# document_processor.py

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import io
import re
from pathlib import Path
from typing import Dict, Optional, Any
import hashlib
import json
from datetime import datetime
import gc
import logging
import numpy as np


class ProductionDocumentProcessor:
    """
    FINAL – Production-grade document processor
    Supports:
    - Digital PDFs
    - Scanned PDFs (OCR)
    - DOCX
    - TXT
    """

    MAX_FILE_SIZE_MB = 20  # Safety guard

    def __init__(self):
        self.logger = self._setup_logging()
        self.cache_dir = Path("cache/document_cache")
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # PUBLIC ENTRY POINT
    # ------------------------------------------------------------------

    def process_document(self, file_path: Path, is_scanned: bool = None) -> Dict[str, Any]:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_path.stat().st_size > self.MAX_FILE_SIZE_MB * 1024 * 1024:
            raise ValueError("File too large to process safely")

        file_hash = self._get_file_hash(file_path)
        cache_file = self.cache_dir / f"{file_hash}.json"

        # Cache check
        if cache_file.exists():
            with open(cache_file, "r", encoding="utf-8") as f:
                cached = json.load(f)
                if self._is_cache_valid(cached):
                    self.logger.info(f"Using cached document: {file_path.name}")
                    return cached

        try:
            ext = file_path.suffix.lower()

            if ext == ".pdf":
                result = self._process_pdf(file_path, is_scanned)
            elif ext in [".docx", ".doc"]:
                result = self._process_word(file_path)
            elif ext == ".txt":
                result = self._process_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")

            # Metadata
            result["metadata"] = {
                "filename": file_path.name,
                "filesize": file_path.stat().st_size,
                "processed_at": datetime.utcnow().isoformat(),
                "file_hash": file_hash,
                "page_count": len(result.get("pages", []))
            }

            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump(result, f, indent=2, default=str)

            return result

        except Exception as e:
            self.logger.error(f"Document processing failed: {e}")
            return self._error_fallback(file_path, str(e))

    # ------------------------------------------------------------------
    # PDF PROCESSING
    # ------------------------------------------------------------------

    def _process_pdf(self, file_path: Path, is_scanned: Optional[bool]) -> Dict[str, Any]:
        if is_scanned is None:
            is_scanned = self._detect_scanned_pdf(file_path)

        return (
            self._process_scanned_pdf(file_path)
            if is_scanned
            else self._process_digital_pdf(file_path)
        )

    def _detect_scanned_pdf(self, file_path: Path) -> bool:
        try:
            with pdfplumber.open(file_path) as pdf:
                char_count = 0
                for page in pdf.pages[:3]:
                    text = page.extract_text()
                    if text:
                        char_count += len(text.strip())
                return char_count < 100
        except Exception:
            return True

    def _process_digital_pdf(self, file_path: Path) -> Dict[str, Any]:
        result = {
            "type": "digital_pdf",
            "pages": [],
            "full_text": "",
            "tables": [],
            "processing_method": "digital_extraction"
        }

        try:
            with fitz.open(file_path) as doc:
                for idx, page in enumerate(doc):
                    text_dict = page.get_text("dict")
                    page_text = self._reconstruct_page_text(text_dict)

                    result["pages"].append({
                        "page_number": idx + 1,
                        "text": page_text,
                        "word_count": len(page_text.split())
                    })

                    result["full_text"] += f"\n\n--- Page {idx + 1} ---\n{page_text}"

                    if idx % 10 == 0:
                        gc.collect()

            # Table extraction (best-effort)
            try:
                with pdfplumber.open(file_path) as pdf:
                    for p_idx, page in enumerate(pdf.pages):
                        for t_idx, table in enumerate(page.extract_tables() or []):
                            cleaned = self._clean_table(table)
                            if cleaned:
                                result["tables"].append({
                                    "page": p_idx + 1,
                                    "table_number": t_idx + 1,
                                    "data": cleaned,
                                    "text_representation": self._table_to_text(cleaned)
                                })
            except Exception:
                pass

            return result

        except Exception as e:
            self.logger.warning(f"Digital PDF failed, falling back to OCR: {e}")
            return self._process_scanned_pdf(file_path)

    def _process_scanned_pdf(self, file_path: Path) -> Dict[str, Any]:
        result = {
            "type": "scanned_pdf",
            "pages": [],
            "full_text": "",
            "processing_method": "ocr"
        }

        try:
            images = convert_from_path(file_path, dpi=300)

            for idx, image in enumerate(images):
                processed = self._preprocess_image(image)

                text = pytesseract.image_to_string(
                    processed,
                    lang="hin+eng",
                    config="--psm 6 --oem 3"
                )

                result["pages"].append({
                    "page_number": idx + 1,
                    "text": text,
                    "word_count": len(text.split()),
                    "ocr_confidence": None
                })

                result["full_text"] += f"\n\n--- Page {idx + 1} ---\n{text}"

                del processed

            return result

        except Exception as e:
            self.logger.error(f"OCR failed: {e}")
            return self._error_fallback(file_path, str(e))

    # ------------------------------------------------------------------
    # WORD & TEXT
    # ------------------------------------------------------------------

    def _process_word(self, file_path: Path) -> Dict[str, Any]:
        doc = Document(file_path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        tables = []
        for t_idx, table in enumerate(doc.tables):
            data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append({
                "table_number": t_idx + 1,
                "data": data,
                "text_representation": self._table_to_text(data)
            })

        return {
            "type": "word_document",
            "pages": [{"page_number": 1, "text": full_text, "word_count": len(full_text.split())}],
            "tables": tables,
            "full_text": full_text,
            "processing_method": "docx_parser"
        }

    def _process_text(self, file_path: Path) -> Dict[str, Any]:
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                return {
                    "type": "text_file",
                    "pages": [{"page_number": 1, "text": text, "word_count": len(text.split())}],
                    "full_text": text,
                    "processing_method": "text_direct"
                }
            except UnicodeDecodeError:
                continue

        with open(file_path, "rb") as f:
            text = f.read().decode("utf-8", errors="replace")

        return {
            "type": "text_file",
            "pages": [{"page_number": 1, "text": text, "word_count": len(text.split())}],
            "full_text": text,
            "processing_method": "text_fallback"
        }

    # ------------------------------------------------------------------
    # HELPERS
    # ------------------------------------------------------------------

    def _reconstruct_page_text(self, page_dict: Dict) -> str:
        blocks = []
        for block in page_dict.get("blocks", []):
            if block.get("type") == 0:
                lines = []
                for line in block.get("lines", []):
                    spans = [span.get("text", "").strip() for span in line.get("spans", []) if span.get("text")]
                    if spans:
                        lines.append(" ".join(spans))
                if lines:
                    blocks.append("\n".join(lines))
        return "\n\n".join(blocks)

    def _clean_table(self, table):
        return [[str(cell).strip() if cell else "" for cell in row] for row in table]

    def _table_to_text(self, table):
        return "\n".join(" | ".join(row) for row in table)

    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        if image.mode != "L":
            image = image.convert("L")
        img = np.array(image)
        img = np.where(img < 128, 0, 255).astype("uint8")
        return Image.fromarray(img)

    def _get_file_hash(self, file_path: Path) -> str:
        hasher = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                hasher.update(chunk)
        return hasher.hexdigest()

    def _is_cache_valid(self, cached: Dict) -> bool:
        try:
            ts = datetime.fromisoformat(cached["metadata"]["processed_at"])
            return (datetime.utcnow() - ts).total_seconds() < 86400
        except Exception:
            return False

    def _error_fallback(self, file_path: Path, error: str) -> Dict[str, Any]:
        return {
            "type": "error",
            "pages": [],
            "full_text": "",
            "error": error,
            "metadata": {
                "filename": file_path.name,
                "processed_at": datetime.utcnow().isoformat(),
                "error_occurred": True
            }
        }

    def _setup_logging(self):
        logger = logging.getLogger("DocumentProcessor")
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                "%(asctime)s - %(name)s - %(levelname)s - %(message)s"
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
            logger.setLevel(logging.INFO)
        return logger
